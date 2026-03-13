"""
FundGate disclosure generator — matches GA sample format exactly.
Font: Arial 9pt throughout. Sig lines via pBdr bottom border.
Supports: LA, FL, GA, KS, MO — 1-signer or 2-signer.
"""
from docx import Document
from docx.shared import Pt, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, datetime

DISCLOSURE_STATES = {
    'LA': {
        'name': 'Louisiana',
        'statute': 'Louisiana Revised Statutes §§9:3573.1–9:3573.8 (Louisiana Commercial Financing Disclosure Law, eff. August 1, 2024)',
        'not_loan': 'This transaction is a purchase and sale of future receivables and is NOT a loan. Amounts charged are NOT interest.',
        'kansas_labels': False,
    },
    'FL': {
        'name': 'Florida',
        'statute': 'Florida Statutes §§559.961–559.9615 (Florida Commercial Financing Disclosure Law, eff. January 1, 2024)',
        'not_loan': 'This transaction is a purchase and sale of future receivables and is NOT a loan. Amounts charged are NOT interest.',
        'kansas_labels': False,
    },
    'GA': {
        'name': 'Georgia',
        'statute': 'Georgia SB 90 (O.C.G.A. § 10-1-393.15 et seq.)',
        'not_loan': 'This transaction is a purchase and sale of future receivables and is NOT a loan. Amounts charged are NOT interest.',
        'kansas_labels': False,
    },
    'KS': {
        'name': 'Kansas',
        'statute': 'Kansas SB 345 — Commercial Financing Disclosure Act (eff. July 1, 2024)',
        'not_loan': 'This transaction is a purchase and sale of future receivables and is NOT a loan. Amounts charged are NOT interest.',
        'kansas_labels': True,
    },
    'MO': {
        'name': 'Missouri',
        'statute': 'Missouri Revised Statutes §427.300 et seq. (Commercial Financing Disclosure Law, eff. February 28, 2025)',
        'not_loan': 'This transaction is a purchase and sale of future receivables and is NOT a loan. Amounts charged are NOT interest.',
        'kansas_labels': False,
    },
}

SZ = 18   # 9pt in half-points (matches sample)
FONT = 'Arial'
W = 10440  # usable width in twips (matches sample: 12240 - 900 - 900)

def _add_border(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'8'); el.set(qn('w:color'),'000000')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _no_border(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'none'); el.set(qn('w:sz'),'0'); el.set(qn('w:color'),'FFFFFF')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}'); el.set(qn('w:w'),str(val)); el.set(qn('w:type'),'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def _col_width(cell, w):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'),str(w)); tcW.set(qn('w:type'),'dxa')
    tcPr.append(tcW)

def _tbl_width(table, w):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0,tblPr)
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'),str(w)); tblW.set(qn('w:type'),'dxa')
    tblPr.append(tblW)

def _vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    vA = OxmlElement('w:vAlign'); vA.set(qn('w:val'),'center'); tcPr.append(vA)

def _spacing(para, before=0, after=60):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'),str(before)); sp.set(qn('w:after'),str(after))
    pPr.append(sp)

def _run(para, text, bold=False, italic=False, size=None, font=None):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    sz = size or SZ
    r.font.size = Pt(sz / 2)
    fn = font or FONT
    r.font.name = fn
    # Set font via XML for all script types (matches sample)
    rPr = r._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    for attr in ('w:ascii','w:cs','w:eastAsia','w:hAnsi'):
        rFonts.set(qn(attr), fn)
    rPr.insert(0, rFonts)
    szEl = OxmlElement('w:sz'); szEl.set(qn('w:val'), str(sz)); rPr.append(szEl)
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(sz)); rPr.append(szCs)
    return r

def _fmt_currency(val):
    try:
        n = float(str(val).replace('$','').replace(',','').replace('%',''))
        return f"${n:,.2f}"
    except:
        return str(val)

def _fmt_date(date_str):
    try:
        for fmt in ('%m/%d/%Y', '%Y-%m-%d'):
            try:
                d = datetime.datetime.strptime(date_str.strip(), fmt)
                return d.strftime('%B %d, %Y')
            except: continue
    except: pass
    return date_str

def _sig_line_para(cell, after=40):
    """Paragraph with bottom border as sig line — matches sample exactly."""
    p = cell.add_paragraph()
    p.clear()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:color'),'000000'); bot.set(qn('w:space'),'1')
    pBdr.append(bot); pPr.append(pBdr)
    _spacing(p, before=0, after=after)
    # Add a space run so the line has height
    _run(p, ' ')
    return p

def build_disclosure_bytes(data):
    state_code = (data.get('State_of_Organization') or '').upper().strip()
    cfg = DISCLOSURE_STATES.get(state_code)
    if not cfg:
        return None

    two_signers   = data.get('twoSigners', False)
    merchant_name = (data.get('Merchant_Legal_Name', '') or '').upper()
    merchant_dba  = (data.get('Merchant_DBA', '') or merchant_name).upper()
    address       = (data.get('Executive_Office_Address', '') or '').upper()
    date_display  = _fmt_date(data.get('Agreement_Date', ''))

    def _n(key):
        try: return float(str(data.get(key,0)).replace('$','').replace(',','').replace('%',''))
        except: return 0.0

    pp       = _n('Purchase_Price')
    pa       = _n('Purchased_Amount')
    orig_pct = _n('Origination_Fee_Percentage')
    orig_amt = round(pp * orig_pct / 100, 2)
    disbursed= round(pp - orig_amt, 2)
    cost     = round(pa - pp, 2)

    pp_fmt   = _fmt_currency(pp)
    pa_fmt   = _fmt_currency(pa)
    orig_fmt = _fmt_currency(orig_amt)
    dis_fmt  = _fmt_currency(disbursed)
    cost_fmt = _fmt_currency(cost)

    spec_pct   = data.get('Specified_Percentage', '')
    weekly_amt = _fmt_currency(_n('Specific_Weekly_Amount'))
    ach_freq   = data.get('ACH_Frequency', 'weekly')

    signer1_name  = (data.get('Owner_Guarantor_1', '') or '').title()
    signer1_title = (data.get('Title', '') or '').title()
    signer2_name  = (data.get('Owner_Guarantor_2', '') or '').title() if two_signers else ''
    signer2_title = (data.get('Title_2', '') or '').title() if two_signers else ''

    kansas = cfg['kansas_labels']

    doc = Document()
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    section = doc.sections[0]
    section.page_width    = Twips(12240)
    section.page_height   = Twips(15840)
    section.top_margin    = Twips(720)
    section.bottom_margin = Twips(720)
    section.left_margin   = Twips(900)
    section.right_margin  = Twips(900)

    # ── Title ──────────────────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(tp, before=0, after=80)
    r = _run(tp, f"{cfg['name'].upper()} COMMERCIAL FINANCING DISCLOSURE", bold=True, size=20)
    r.font.underline = True

    # ── Date ───────────────────────────────────────────────────────────────────
    dp = doc.add_paragraph()
    _spacing(dp, before=0, after=80)
    _run(dp, 'Disclosure Date: ')
    _run(dp, date_display, bold=True)

    # ── Header table ───────────────────────────────────────────────────────────
    ht = doc.add_table(rows=2, cols=2)
    _tbl_width(ht, W)
    for row in ht.rows:
        for cell in row.cells:
            _col_width(cell, W//2); _add_border(cell)
            _cell_margins(cell, top=100, bottom=100, left=120, right=120)

    lc = ht.cell(0,0); lc.paragraphs[0].clear()
    _run(lc.paragraphs[0], 'Recipient: ', bold=True); _run(lc.paragraphs[0], merchant_name, bold=True)
    _spacing(lc.paragraphs[0], after=40)
    p2 = lc.add_paragraph(); _run(p2,'DBA: ',bold=True); _run(p2,merchant_dba,bold=True); _spacing(p2,after=40)
    p3 = lc.add_paragraph(); _run(p3,'Address: ',bold=True); _run(p3,address,bold=True); _spacing(p3,after=0)

    rc = ht.cell(0,1); rc.paragraphs[0].clear()
    _run(rc.paragraphs[0], 'Provider', bold=True); _spacing(rc.paragraphs[0], after=40)
    for label, val in [('Name: ','FundGate LLC'),('Address: ','1202 Avenue U, Suite 1175, Brooklyn NY 11229'),
                       ('Phone: ','929-355-8918'),('Email: ','admin@fundgatellc.com')]:
        px = rc.add_paragraph()
        _run(px, label, bold=True); _run(px, val, bold=True)
        _spacing(px, after=0 if label=='Email: ' else 40)

    merged = ht.cell(1,0).merge(ht.cell(1,1))
    _add_border(merged); _cell_margins(merged,top=80,bottom=80,left=120,right=120)
    merged.paragraphs[0].clear()
    _run(merged.paragraphs[0],
         f'This Commercial Financing Disclosure is being provided to the Recipient ("you") by the '
         f'Provider ("we" or "us") as required by {cfg["statute"]} and is dated as of the Disclosure Date.',
         italic=True)
    _spacing(merged.paragraphs[0], after=0)

    # ── Amounts table ──────────────────────────────────────────────────────────
    if kansas:
        rows_spec = [
            ('1.  Total Amount of Funds Provided', pp_fmt),
            None,
            ('3.  Total of Payments', pa_fmt),
            ('4.  Total Dollar Cost of Financing', cost_fmt),
        ]
        r2_label = '2.  Total Amount of Funds Disbursed'
        r2_right = dis_fmt
    else:
        rows_spec = [
            ('1.  Total Amount of Funding Provided', pp_fmt),
            None,
            ('3.  Total Amount of Funds Disbursed (1 minus 2)', dis_fmt),
            ('4.  Total Amount to be Paid to Us', pa_fmt),
            ('5.  Total Dollar Cost (4 minus 1)', cost_fmt),
        ]
        r2_label = '2.  Amounts Deducted from Funding Provided'
        r2_right = orig_fmt

    tt = doc.add_table(rows=0, cols=2)
    _tbl_width(tt, W)

    def bold_row(label, amount):
        row = tt.add_row(); lc2, rc2 = row.cells
        _col_width(lc2, 7880); _col_width(rc2, 2560)
        _add_border(lc2); _add_border(rc2)
        _cell_margins(lc2,top=60,bottom=60,left=120,right=80)
        _cell_margins(rc2,top=60,bottom=60,left=80,right=120)
        lc2.paragraphs[0].clear(); _run(lc2.paragraphs[0], label, bold=True); _spacing(lc2.paragraphs[0],after=0)
        rc2.paragraphs[0].clear(); rc2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(rc2.paragraphs[0], amount, bold=True); _spacing(rc2.paragraphs[0],after=0)
        _vcenter(rc2)

    for r in rows_spec:
        if r is None:
            row2 = tt.add_row(); lc2, rc2 = row2.cells
            _col_width(lc2,7880); _col_width(rc2,2560)
            _add_border(lc2); _add_border(rc2)
            _cell_margins(lc2,top=60,bottom=60,left=120,right=80)
            _cell_margins(rc2,top=60,bottom=60,left=80,right=120)
            lc2.paragraphs[0].clear()
            _run(lc2.paragraphs[0], r2_label, bold=True); _spacing(lc2.paragraphs[0],after=40)
            for line in [
                f'   Fees deducted or withheld at disbursement .......................................  {orig_fmt}',
                '   Amount deducted for prior balance paid to us ...................................  $0.00',
                '   Amount deducted and paid to third parties on your behalf .......................  $0.00',
            ]:
                px2 = lc2.add_paragraph(); _run(px2, line); _spacing(px2,after=0)
            rc2.paragraphs[0].clear(); rc2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _run(rc2.paragraphs[0], r2_right, bold=True); _spacing(rc2.paragraphs[0],after=0)
            _vcenter(rc2)
        else:
            bold_row(r[0], r[1])

    # ── Payment / prepayment table ─────────────────────────────────────────────
    freq_word = 'Business Week' if 'week' in ach_freq.lower() else 'Business Day'
    bt = doc.add_table(rows=0, cols=2)
    _tbl_width(bt, W)

    def wide_row(label, build_fn):
        row = bt.add_row(); lc3, rc3 = row.cells
        _col_width(lc3,2800); _col_width(rc3,7640)
        _add_border(lc3); _add_border(rc3)
        _cell_margins(lc3,top=80,bottom=80,left=120,right=80)
        _cell_margins(rc3,top=80,bottom=80,left=80,right=120)
        _vcenter(lc3)
        lc3.paragraphs[0].clear(); _run(lc3.paragraphs[0], label, bold=True); _spacing(lc3.paragraphs[0],after=0)
        rc3.paragraphs[0].clear(); _spacing(rc3.paragraphs[0],after=0)
        build_fn(rc3)

    def build_payment(cell):
        p0 = cell.paragraphs[0]
        _run(p0, 'We will collect the Total Amount to be Paid to Us by debiting your business bank '
             'account in periodic installments or "payments" that will occur with the following frequency:')
        _spacing(p0,after=40)
        p1 = cell.add_paragraph()
        _run(p1, f'☒ Every {freq_word}', bold=True)
        _run(p1, '  (i.e., one debit per week on a designated business day, excluding bank holidays. '
             'Payments scheduled for a bank holiday will be debited the next business day with the regular payment)')
        _spacing(p1,after=40)
        p2 = cell.add_paragraph()
        _run(p2, 'The initial payment will be ')
        _run(p2, weekly_amt, bold=True)
        _run(p2, '. We based your initial payment on ')
        _run(p2, str(spec_pct), bold=True)
        _run(p2, ' of your estimated sales revenue. For details on your right to adjust any payment amount, '
             'see Section 3 of your Purchase Agreement.')
        _spacing(p2,after=0)

    def build_prepay(cell):
        p0 = cell.paragraphs[0]
        _run(p0, 'If you pay off the financing faster than required, you may pay a reduced amount per '
             'the Addendum to Merchant Cash Advance Agreement dated ')
        _run(p0, date_display)
        _run(p0, ', which sets forth the contractual rights of the parties related to prepayment. '
             'No additional fees will be charged for prepayment.')
        _spacing(p0,after=0)

    pay_label    = 'Estimated Payments' if kansas else 'Manner, frequency, and amount of each payment'
    wide_row(pay_label, build_payment)
    wide_row('Description of Prepayment Policies', build_prepay)

    # ── Acknowledgment ─────────────────────────────────────────────────────────
    ack = doc.add_paragraph()
    _run(ack, 'By signing below, you acknowledge that you have received a copy of this disclosure form.')
    _spacing(ack, before=80, after=80)

    # ── Signature table — matches sample exactly ───────────────────────────────
    # Cols: sig(5100) | spacer(600) | date(4740) — from sample
    st = doc.add_table(rows=1, cols=3)
    _tbl_width(st, W)
    lsig, sp, rsig = st.rows[0].cells
    _col_width(lsig, 5100); _col_width(sp, 600); _col_width(rsig, 4740)
    for c in [lsig, sp, rsig]:
        _no_border(c); _cell_margins(c)
        c.paragraphs[0].clear()

    def _add_signer(sig_cell, date_cell, name, title, spacer=False):
        if spacer:
            sp2 = sig_cell.add_paragraph(); _spacing(sp2, before=60, after=60)
            sp3 = date_cell.add_paragraph(); _spacing(sp3, before=60, after=60)
        # Sig line
        _sig_line_para(sig_cell, after=40)
        lp = sig_cell.add_paragraph()
        label = f'Recipient Signature \u2014 {name}, {title}' if title else f'Recipient Signature \u2014 {name}'
        _run(lp, label); _spacing(lp, before=0, after=40 if spacer else 0)
        # Date line
        _sig_line_para(date_cell, after=40)
        dp2 = date_cell.add_paragraph(); _run(dp2, 'Date'); _spacing(dp2, before=0, after=40 if spacer else 0)

    # Remove initial empty paragraphs
    for c in [lsig, sp, rsig]:
        for p in c.paragraphs:
            p._element.getparent().remove(p._element)

    _add_signer(lsig, rsig, signer1_name, signer1_title)
    if two_signers and signer2_name:
        _add_signer(lsig, rsig, signer2_name, signer2_title, spacer=True)

    # ── Statute footer ─────────────────────────────────────────────────────────
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fp, f"Pursuant to {cfg['statute']}. {cfg['not_loan']}", italic=True, size=16)
    _spacing(fp, before=80, after=0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
